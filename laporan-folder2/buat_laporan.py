from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Margin halaman ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph spacing ──
def set_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

# ── Helper: heading berwarna ──
def add_heading(doc, text, level=1, color=RGBColor(0x1a,0x1a,0x1a)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
    set_spacing(h, before=120, after=60)
    return h

# ══════════════════════════════════════════
# COVER / HALAMAN JUDUL
# ══════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(cover, before=600, after=0)

r = cover.add_run('LAPORAN PRAKTIKUM\nPEMROGRAMAN WEB')
r.font.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x5C,0x3D,0x2E)

doc.add_paragraph()
doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('Arsitektur MVC & Konfigurasi Database\nAplikasi E-Commerce Furniture "FurniNest"')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x55,0x55,0x55)
set_spacing(sub, before=0, after=0)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Tabel identitas
tbl = doc.add_table(rows=3, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

data = [
    ('Nama',          'Muhamad Daffa Nasution'),
    ('NIM',           '20240801289'),
    ('Mata Kuliah',   'Pemrograman Web'),
]
for i, (label, val) in enumerate(data):
    tbl.rows[i].cells[0].text = label
    tbl.rows[i].cells[1].text = val
    for cell in tbl.rows[i].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(12)
                if cell == tbl.rows[i].cells[0]:
                    run.font.bold = True

# Warna header kolom label
for i in range(3):
    tc = tbl.rows[i].cells[0]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F0E8')
    tcPr.append(shd)

doc.add_page_break()

# ══════════════════════════════════════════
# BAB 1 — PENDAHULUAN
# ══════════════════════════════════════════
add_heading(doc, 'BAB I — Pendahuluan', level=1, color=RGBColor(0x5C,0x3D,0x2E))

intro = doc.add_paragraph(
    'Laporan ini menjelaskan implementasi arsitektur MVC (Model-View-Controller) '
    'pada proyek aplikasi e-commerce furniture berbasis web yang dibangun menggunakan '
    'framework Laravel (PHP). Arsitektur MVC memisahkan logika aplikasi menjadi tiga '
    'komponen utama sehingga kode lebih terstruktur, mudah dipelihara, dan dapat '
    'dikembangkan secara mandiri oleh masing-masing lapisan.'
)
intro.style.font.size = Pt(12)
set_spacing(intro, before=0, after=120)

# ══════════════════════════════════════════
# BAB 2 — PENJELASAN MVC
# ══════════════════════════════════════════
add_heading(doc, 'BAB II — Struktur MVC Proyek FurniNest', level=1, color=RGBColor(0x5C,0x3D,0x2E))

# Gambar MVC (screenshot)
ss_path = r'c:\web-copy-php\laporan-folder2\ss_mvc.png'
if os.path.exists(ss_path):
    caption_before = doc.add_paragraph('Gambar 1. Visualisasi Struktur Folder MVC dan Konfigurasi Database')
    caption_before.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_before.runs:
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x77,0x77,0x77)
    set_spacing(caption_before, before=0, after=80)
    
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pic = pic_para.add_run()
    run_pic.add_picture(ss_path, width=Cm(15))
    set_spacing(pic_para, before=0, after=120)

doc.add_paragraph()

# ── 2.1 MODEL ──
add_heading(doc, '2.1  Model  (app/Models/)', level=2, color=RGBColor(0x7C,0x3A,0xED))

p_model = doc.add_paragraph(
    'Model merupakan lapisan yang bertanggung jawab atas seluruh interaksi dengan database. '
    'Setiap file Model merepresentasikan satu tabel dalam database dan dikelola menggunakan '
    'Eloquent ORM bawaan Laravel. Model juga mendefinisikan relasi antar entitas, seperti '
    'hubungan antara User dengan OrderHistory, atau Product dengan CartItem. '
    'Dengan adanya Model, logika akses data terpusat dan tidak tersebar di seluruh bagian aplikasi.'
)
p_model.style.font.size = Pt(12)
set_spacing(p_model, before=0, after=120)

# ── 2.2 VIEW ──
add_heading(doc, '2.2  View  (resources/views/)', level=2, color=RGBColor(0x0E,0xA5,0xE9))

p_view = doc.add_paragraph(
    'View adalah lapisan yang bertugas menampilkan antarmuka pengguna (UI). Pada proyek ini, '
    'View menggunakan Blade Template Engine milik Laravel yang memungkinkan penulisan '
    'kode PHP langsung di dalam HTML dengan sintaks yang bersih dan mudah dibaca. '
    'Folder admin/ berisi tampilan khusus panel administrator, sedangkan folder pages/ '
    'berisi halaman-halaman yang diakses oleh pelanggan seperti beranda, detail produk, '
    'checkout, login, dan register. View hanya bertugas menampilkan data — tidak mengandung logika bisnis.'
)
p_view.style.font.size = Pt(12)
set_spacing(p_view, before=0, after=120)

# ── 2.3 CONTROLLER ──
add_heading(doc, '2.3  Controller  (app/Http/Controllers/)', level=2, color=RGBColor(0xF5,0x9E,0x0B))

p_ctrl = doc.add_paragraph(
    'Controller adalah lapisan penghubung antara Model dan View. Ketika pengguna mengakses '
    'sebuah URL, Laravel meneruskan request tersebut ke Controller yang sesuai. Controller '
    'kemudian memanggil Model untuk mengambil atau memanipulasi data di database, lalu '
    'mengirimkan hasilnya ke View untuk ditampilkan. Sebagai contoh, AdminDashboardController '
    'mengambil data total penjualan, total pesanan, total produk, dan daftar pesanan terbaru '
    'dari database, kemudian meneruskan semua data tersebut ke halaman dashboard admin.'
)
p_ctrl.style.font.size = Pt(12)
set_spacing(p_ctrl, before=0, after=120)

# ══════════════════════════════════════════
# BAB 3 — DATABASE
# ══════════════════════════════════════════
add_heading(doc, 'BAB III — Konfigurasi Database', level=1, color=RGBColor(0x5C,0x3D,0x2E))

add_heading(doc, '3.1  Konfigurasi Koneksi  (config/database.php)', level=2, color=RGBColor(0x10,0xB9,0x81))

p_db = doc.add_paragraph(
    'Konfigurasi database pada proyek FurniNest menggunakan MySQL sebagai sistem manajemen '
    'basis data (DBMS). Koneksi dikonfigurasi melalui file config/database.php yang '
    'membaca nilai sensitif (host, port, nama database, username, password) dari file .env '
    'sehingga konfigurasi aman dan mudah diganti sesuai lingkungan (development/production). '
    'Proyek ini menggunakan port 3307 dan nama database "furninest".'
)
p_db.style.font.size = Pt(12)
set_spacing(p_db, before=0, after=120)

# Tabel struktur database
add_heading(doc, '3.2  Tabel-Tabel Utama Database', level=2, color=RGBColor(0x10,0xB9,0x81))

p_tbl_intro = doc.add_paragraph('Berikut adalah tabel-tabel utama yang digunakan dalam aplikasi FurniNest:')
p_tbl_intro.style.font.size = Pt(12)
set_spacing(p_tbl_intro, before=0, after=80)

db_table = doc.add_table(rows=1, cols=3)
db_table.style = 'Table Grid'
db_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr = db_table.rows[0].cells
hdr[0].text = 'Nama Tabel'
hdr[1].text = 'Model'
hdr[2].text = 'Fungsi'
for cell in hdr:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(11)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '5C3D2E')
    tcPr.append(shd)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

db_rows = [
    ('users',          'User.php',         'Menyimpan data akun pelanggan dan admin'),
    ('products',       'Product.php',      'Menyimpan data produk furniture (nama, harga, stok, gambar)'),
    ('order_histories','OrderHistory.php', 'Menyimpan riwayat pesanan pelanggan beserta status'),
    ('order_items',    'OrderItem.php',    'Detail produk dalam setiap pesanan'),
    ('cart_items',     'CartItem.php',     'Data keranjang belanja sementara pelanggan'),
    ('addresses',      'Address.php',      'Alamat pengiriman milik pelanggan'),
    ('wishlists',      'Wishlist.php',     'Daftar produk favorit yang disimpan pelanggan'),
]

for i, (tbl_name, model, func) in enumerate(db_rows):
    row = db_table.add_row()
    row.cells[0].text = tbl_name
    row.cells[1].text = model
    row.cells[2].text = func
    fill = 'FAFAFA' if i % 2 == 0 else 'FFFFFF'
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

doc.add_paragraph()

# ══════════════════════════════════════════
# PENUTUP
# ══════════════════════════════════════════
add_heading(doc, 'BAB IV — Penutup', level=1, color=RGBColor(0x5C,0x3D,0x2E))

penutup = doc.add_paragraph(
    'Arsitektur MVC yang diterapkan pada proyek FurniNest memberikan struktur kode yang '
    'terorganisir dan mudah dikembangkan. Pemisahan antara Model, View, dan Controller '
    'menjadikan setiap lapisan memiliki tanggung jawab yang jelas — Model mengelola data, '
    'Controller memproses logika, dan View menampilkan hasil — sehingga pengembangan '
    'dan pemeliharaan aplikasi menjadi lebih efisien dan sistematis.'
)
penutup.style.font.size = Pt(12)
set_spacing(penutup, before=0, after=120)

# ── Simpan ──
out_path = r'c:\web-copy-php\laporan-folder2\Laporan_Pemrograman_Web_Daffa.docx'
doc.save(out_path)
print(f'Berhasil! File disimpan di: {out_path}')
